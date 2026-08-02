"""Build paper/draft.docx from paper/sections/*.md.

Markdown is the source of truth; this script regenerates the Word file. Never edit
draft.docx by hand — edits there are lost on the next build.

Handles the markdown subset actually used in the sections: H1/H2 headings, paragraphs,
**bold**, *italic*, `code`, and figure references. A paragraph that mentions
`fig_<name>.png` gets that figure inserted after it with a numbered caption drawn from
figures/FIGURES.md. HTML comments (drafting scaffolding) are stripped. Remaining
TODO(...) markers are rendered in red so they are impossible to miss in review.

Usage:  python paper/build_docx.py [-o draft.docx]
"""
from __future__ import annotations

import argparse
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
SECTIONS = os.path.join(HERE, "sections")
FIGDIR = os.path.join(HERE, "figures")

TITLE = ("Measuring Sustained, Intent-Grounded Deception Among Competing "
         "Frontier Models in an Open Survival Economy")
SUBTITLE = "Working draft — all results are single-seed observations"

# Figure captions, keyed by filename. Kept short; full provenance is in FIGURES.md.
CAPTIONS = {
    "fig_leaderboard_9model.png":
        "Deception rate by model in the 335-turn ten-model run, exposure-normalised, "
        "with turns survived shown alongside so the ~9x exposure spread is visible "
        "rather than buried.",
    "fig_coherence_335t.png":
        "Coherence against a permutation null: campaign raster, target selectivity, "
        "longest resumed thread, and rate by 50-turn bucket. Most raw coherence is "
        "explained by the null.",
    "fig_model_not_seat.png":
        "Two instances of each model, differing in specialty, rivals, and fate, land in "
        "the same deception band; different models do not. Within-game control.",
    "fig_deception_pay_flagship.png":
        "Deception rate against final wealth in the flagship game. The apex winner was "
        "among the least deceptive; ring marks elimination.",
    "fig_condition_contrast.png":
        "Deception rate by model and condition. The honesty instruction lowers deception; "
        "explicit permission barely raises it.",
    "fig_type_mix_condition.png":
        "Deception type mix by condition: instruction changes the kind of deception, not "
        "only its frequency.",
    "fig_type_by_model.png":
        "Per-model deception signature: models differ in which deception types they use.",
    "fig_campaigns.png":
        "Illustrative targeted campaigns over turns, coloured by target.",
}

FIG_RE = re.compile(r"fig_[a-z0-9_]+\.png")
# bold | italic | code | TODO marker
TOKEN_RE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*\n]+?\*(?!\*)|`[^`]+`|TODO\([a-z]+\))")


def strip_comments(text: str) -> str:
    return re.sub(r"<!--.*?-->", "", text, flags=re.S)


def add_runs(par, text: str) -> None:
    """Add markdown-inline-formatted runs to a docx paragraph."""
    text = text.replace(r"\$", "$").replace("\\_", "_")
    for tok in TOKEN_RE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            inner = tok[1:-1]
            if inner.startswith("TODO("):  # backtick-wrapped TODO: keep it loud, not code
                r = par.add_run(inner)
                r.bold = True
                r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            else:
                r = par.add_run(inner)
                r.font.name = "Consolas"
                r.font.size = Pt(9.5)
        elif tok.startswith("TODO("):
            r = par.add_run(tok)
            r.bold = True
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            par.add_run(tok[1:-1]).italic = True
        else:
            par.add_run(tok)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("-o", "--out", default=os.path.join(HERE, "draft.docx"))
    args = ap.parse_args()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(1)
        s.left_margin = s.right_margin = Inches(1)

    # --- title block ---
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(16)
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run(SUBTITLE)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    files = sorted(f for f in os.listdir(SECTIONS) if f.endswith(".md"))
    fig_no = 0
    seen_figs: set[str] = set()
    section_no = 0

    for fname in files:
        raw = strip_comments(open(os.path.join(SECTIONS, fname), encoding="utf-8").read())
        blocks = [b.strip() for b in raw.split("\n\n") if b.strip()]
        is_abstract = fname.startswith("00")

        for block in blocks:
            lines = block.split("\n")
            first = lines[0]

            if first.startswith("## "):
                doc.add_heading(first[3:].strip(), level=2)
                rest = " ".join(x.strip() for x in lines[1:]).strip()
                if rest:
                    add_runs(doc.add_paragraph(), rest)
                continue

            if first.startswith("# "):
                title = first[2:].strip()
                if not is_abstract:
                    section_no += 1
                    title = f"{section_no}. {title}"
                doc.add_heading(title, level=1)
                rest = " ".join(x.strip() for x in lines[1:]).strip()
                if rest:
                    add_runs(doc.add_paragraph(), rest)
                continue

            # body paragraph: join wrapped lines
            text = " ".join(x.strip() for x in lines).strip()
            if not text:
                continue
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(par, text)

            # insert any figure this paragraph references (first mention only)
            for fig in FIG_RE.findall(text):
                if fig in seen_figs:
                    continue
                path = os.path.join(FIGDIR, fig)
                if not os.path.exists(path):
                    continue
                seen_figs.add(fig)
                fig_no += 1
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=Inches(5.6))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cap.add_run(f"Figure {fig_no}. {CAPTIONS.get(fig, '')}")
                cr.font.size = Pt(9)
                cr.italic = True

    doc.save(args.out)
    words = sum(
        len(strip_comments(open(os.path.join(SECTIONS, f), encoding="utf-8").read()).split())
        for f in files
    )
    print(f"wrote {args.out}")
    print(f"  sections: {len(files)}  figures placed: {fig_no}  ~{words} words")
    unplaced = sorted(set(CAPTIONS) - seen_figs)
    if unplaced:
        print("  not placed (never referenced in text):", ", ".join(unplaced))


if __name__ == "__main__":
    main()
